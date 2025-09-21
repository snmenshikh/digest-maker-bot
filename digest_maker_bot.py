import os
import logging
import asyncio
from datetime import datetime, timedelta, timezone

import pandas as pd
import requests
from bs4 import BeautifulSoup

from aiogram import Bot, Dispatcher, types
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.filters import Command

import nltk
from nltk.tokenize import sent_tokenize

from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
import docx.opc.constants

# ---------- Логи ----------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("digest_maker_bot")

# ---------- FSM состояния ----------
class DigestStates(StatesGroup):
    WAITING_FOR_EXCEL = State()
    WAITING_FOR_INTERVAL = State()
    WAITING_FOR_KEYWORDS = State()

# ---------- Токен ----------
BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN не найден. Установите переменную окружения в Portainer.")

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())

# ---------- NLTK ----------
nltk.download("punkt")

# ---------- DOCX: гиперссылки ----------
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# ---------- Очистка текста ----------
def clean_text(text: str) -> str:
    return text.replace("\x00", "").replace("\u0000", "").strip()

# ---------- Суммаризация ----------
def summarize_text_extractively(text, keywords, max_sentences=3):
    sentences = [s.strip() for s in sent_tokenize(text) if s.strip()]
    if not sentences:
        return ""
    if keywords:
        # приводим ключевые слова к нижнему регистру
        keywords = [k.lower() for k in keywords]
        scored = [s for s in sentences if any(k in s.lower() for k in keywords)]
    else:
        scored = sentences
    return " ".join(scored[:max_sentences])

# ---------- Парсинг канала ----------
def parse_channel(url, keywords, start_dt):
    try:
        html = requests.get(url, timeout=15).text
        soup = BeautifulSoup(html, "html.parser")
        posts = soup.find_all("div", class_="tgme_widget_message_wrap")

        items = []
        seen = set()
        for post in posts:
            dt_node = post.find("time")
            dt = None
            if dt_node and dt_node.has_attr("datetime"):
                dt = datetime.fromisoformat(dt_node["datetime"].replace("Z", "+00:00"))
                if dt < start_dt:
                    continue

            text_node = post.find("div", class_="tgme_widget_message_text")
            if not text_node:
                continue
            text = clean_text(text_node.get_text(" "))

            post_id = post.get("data-post")
            post_url = None
            if post_id and "/" in post_id:
                channel_username, msg_id = post_id.split("/")
                post_url = f"https://t.me/{channel_username}/{msg_id}"

            summary = summarize_text_extractively(text, keywords, 3)
            if not summary:
                continue  # пропускаем посты без совпадений по ключевым словам

            key = (dt, summary)
            if key in seen:
                continue
            seen.add(key)

            items.append({
                "dt": dt,
                "summary": summary,
                "original": text,
                "post_url": post_url,
            })
        return items
    except Exception as e:
        logger.error(f"Ошибка парсинга канала {url}: {e}")
        return []

# ---------- Генерация .docx ----------
def build_docx_digest(user_id, channels, keywords, interval_days):
    doc = Document()
    doc.add_heading("Дайджест по Telegram-каналам", level=1)

    local_tz = timezone(timedelta(hours=5))
    doc.add_paragraph(f"Сформирован: {datetime.now(local_tz).strftime('%Y-%m-%d %H:%M:%S')}")

    start_dt = datetime.now(timezone.utc) - timedelta(days=interval_days)

    for ch_name, url in channels:
        items = parse_channel(url, keywords, start_dt)
        if not items:
            continue  # пропускаем каналы без публикаций

        hdr = doc.add_heading(ch_name, level=2)
        hdr.style.font.size = Pt(13)
        doc.add_paragraph(f"Источник: {ch_name} ({url})")

        for it in items:
            dt_str = it["dt"].astimezone(local_tz).strftime("%Y-%m-%d %H:%M") if it["dt"] else "дата не распознана"
            doc.add_paragraph(f"Дата публикации: {dt_str}")

            doc.add_paragraph(it["summary"])

            if it.get("post_url"):
                add_hyperlink(doc.add_paragraph(), "🔗 Открыть оригинал", it["post_url"])

            doc.add_paragraph("---")

    fname = f"digest_{user_id}_{int(datetime.now().timestamp())}.docx"
    path = os.path.join(os.getcwd(), fname)
    doc.save(path)
    return path

# ---------- Обработчики ----------
async def on_start(message: Message, state: FSMContext):
    await state.clear()
    await message.answer(
        "Привет! Я подготовлю дайджест по Telegram-каналам.\n"
        "Пришлите Excel-файл со списком каналов:\n"
        "• столбец A — имя канала\n"
        "• столбец B — ссылка вида https://t.me/<slug>\n\n"
        "Поддерживаются *.xls и *.xlsx."
    )
    await state.set_state(DigestStates.WAITING_FOR_EXCEL)

async def on_excel(message: Message, state: FSMContext):
    if not message.document:
        await message.answer("Пожалуйста, пришлите файл Excel.")
        return

    file_path = f"temp_{message.from_user.id}.xlsx"
    await bot.download(message.document, destination=file_path)

    try:
        df = pd.read_excel(file_path, engine="openpyxl")
        if df.shape[1] < 2:
            raise ValueError("Файл должен содержать минимум 2 столбца: имя и ссылку.")
        channels = df.iloc[:, :2].dropna().values.tolist()
    except Exception as e:
        await message.answer(f"Ошибка чтения Excel: {e}")
        return
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

    await state.update_data(channels=channels)

    kb = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="Сутки"), KeyboardButton(text="Неделя"), KeyboardButton(text="Месяц")]],
        resize_keyboard=True,
        one_time_keyboard=True,
    )
    await message.answer("Выберите интервал времени:", reply_markup=kb)
    await state.set_state(DigestStates.WAITING_FOR_INTERVAL)

async def on_interval(message: Message, state: FSMContext):
    intervals = {"Сутки": 1, "Неделя": 7, "Месяц": 30}
    interval_days = intervals.get(message.text)
    if not interval_days:
        await message.answer("Пожалуйста, выберите интервал из кнопок.")
        return

    await state.update_data(interval_days=interval_days)
    await message.answer("Введите ключевые слова (через запятую):")
    await state.set_state(DigestStates.WAITING_FOR_KEYWORDS)

async def on_keywords(message: Message, state: FSMContext):
    data = await state.get_data()
    channels = data["channels"]
    interval_days = data["interval_days"]
    keywords = [k.strip().lower() for k in message.text.split(",") if k.strip()]

    try:
        out_path = build_docx_digest(message.from_user.id, channels, keywords, interval_days)
        await message.answer_document(types.FSInputFile(out_path))
        os.remove(out_path)
    except Exception as e:
        logger.error(f"Unexpected processing error: {e}", exc_info=True)
        await message.answer("Ошибка при обработке. Попробуйте снова.")

    await state.clear()

# ---------- Регистрация обработчиков ----------
dp.message.register(on_start, Command(commands=["start"]))
dp.message.register(on_excel, DigestStates.WAITING_FOR_EXCEL)
dp.message.register(on_interval, DigestStates.WAITING_FOR_INTERVAL)
dp.message.register(on_keywords, DigestStates.WAITING_FOR_KEYWORDS)

# ---------- Main ----------
async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())